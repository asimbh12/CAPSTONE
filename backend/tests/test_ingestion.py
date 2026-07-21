from io import BytesIO

from docx import Document as DocxDocument
from fastapi.testclient import TestClient

from app.schemas.ingestion import (
    CareerExtractionProposal,
    ProposedAsset,
    ProposedProfile,
    PublicProfileSource,
)
from app.services.ai import DeterministicCareerExtractor
from app.services.documents import _extract_text
from app.services.ingestion import (
    discover_linked_profile_sources,
    expand_public_profile_sources,
    extract_google_scholar_rows,
    is_thin_scholar_continuation,
    merge_source_proposals,
)


def test_docx_extraction_includes_tables() -> None:
    document = DocxDocument()
    document.add_paragraph("Public career profile")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Director"
    table.cell(0, 1).text = "2020 – present"
    content = BytesIO()
    document.save(content)

    extracted = _extract_text(".docx", content.getvalue())

    assert "Director | 2020 – present" in extracted


def test_offline_extraction_does_not_treat_url_digits_as_a_year() -> None:
    proposal = DeterministicCareerExtractor().extract(
        "Example Professional\nLinkedIn: https://example.org/profile-1907", "CV"
    )
    assert proposal.assets == []


def test_deakin_expert_profile_expands_to_all_public_sections() -> None:
    sources = expand_public_profile_sources(
        "https://experts.deakin.edu.au/150-asim-bhatti/publications"
    )

    assert [source.source_type for source in sources] == [
        "profile",
        "research_outputs",
        "research_grants",
        "professional_activities",
        "teaching_supervision",
    ]
    assert sources[0].url == "https://experts.deakin.edu.au/150-asim-bhatti"
    assert sources[-1].url.endswith("/teaching")


def test_non_deakin_profile_is_not_expanded() -> None:
    sources = expand_public_profile_sources("https://orcid.org/0000-0001-6876-1437")
    assert len(sources) == 1
    assert sources[0].url == "https://orcid.org/0000-0001-6876-1437"


def test_google_scholar_profile_expands_to_hundred_result_pages() -> None:
    sources = expand_public_profile_sources(
        "https://scholar.google.com/citations?user=abc123&hl=en"
    )

    assert len(sources) == 10
    assert all(source.source_type == "google_scholar" for source in sources)
    assert "cstart=0" in sources[0].url and "pagesize=100" in sources[0].url
    assert "cstart=900" in sources[-1].url


def test_google_scholar_thin_continuation_stops_before_ai_extraction() -> None:
    continuation = PublicProfileSource(
        url="https://scholar.google.com/citations?user=abc123&cstart=200&pagesize=100",
        source_type="google_scholar",
    )

    assert is_thin_scholar_continuation(continuation, "empty page shell" * 50)
    assert not is_thin_scholar_continuation(continuation, "publication content" * 500)


def test_google_scholar_rows_are_extracted_without_llm_summarisation() -> None:
    html = """
    <table>
      <tr class="gsc_a_tr"><td><a class="gsc_a_at">First publication</a>
      <div class="gs_gray">A Author, B Author</div><div class="gs_gray">Journal 12</div></td>
      <td class="gsc_a_y"><span>2024</span></td></tr>
      <tr class="gsc_a_tr"><td><a class="gsc_a_at">Second publication</a>
      <div class="gs_gray">A Author</div></td><td class="gsc_a_y"><span>2023</span></td></tr>
    </table>
    """

    proposal = extract_google_scholar_rows(html, "Asim Bhatti - Google Scholar")

    assert proposal is not None
    assert proposal.profile.name == "Asim Bhatti"
    assert [asset.title for asset in proposal.assets] == ["First publication", "Second publication"]
    assert proposal.assets[0].start_date.isoformat() == "2024-01-01"


def test_link_discovery_keeps_relevant_same_site_pages_only() -> None:
    source = PublicProfileSource(
        url="https://profiles.example.edu/people/jane-doe", source_type="institutional_profile"
    )
    html = """
    <a href="/people/jane-doe/publications?page=2">More publications</a>
    <a href="/people/jane-doe/grants">Research grants</a>
    <a href="https://unrelated.example.org/jane">External profile</a>
    <a href="/privacy">Privacy</a>
    """

    sources = discover_linked_profile_sources(source, html)

    assert [item.url for item in sources] == [
        source.url,
        "https://profiles.example.edu/people/jane-doe/publications?page=2",
        "https://profiles.example.edu/people/jane-doe/grants",
    ]


def test_document_ingestion_reviews_and_populates_without_overwriting_user_facts(
    client: TestClient,
) -> None:
    client.put(
        "/api/profile",
        json={
            "name": "User-authored name",
            "current_title": "",
            "current_organisation": "",
            "career_mission": "User-authored mission",
            "career_narrative": "",
        },
    )
    document = (
        b"Example Professional\nDirector, Example Institute 2020 - present\nLed public programs."
    )
    response = client.post(
        "/api/ingestions/documents",
        files={"file": ("career.txt", document, "text/plain")},
        data={
            "ai_handling_policy": "local_only",
            "confirmed_public_information": "true",
        },
    )
    assert response.status_code == 201
    run = response.json()
    assert run["provider"] == "deterministic"
    assert len(run["proposal"]["assets"]) == 1
    run["proposal"]["profile"]["current_title"] = "Proposed director"

    applied = client.post(f"/api/ingestions/{run['id']}/apply", json={"proposal": run["proposal"]})
    assert applied.status_code == 200
    assert applied.json()["assets_created"] == 1
    profile = client.get("/api/profile").json()
    assert profile["name"] == "User-authored name"
    assert profile["career_mission"] == "User-authored mission"
    assert profile["current_title"] == "Proposed director"
    assert client.get("/api/timeline").json()[0]["start_date"] == "2020-01-01"


def test_ingestion_rejects_unconfirmed_document(client: TestClient) -> None:
    response = client.post(
        "/api/ingestions/documents",
        files={"file": ("career.txt", b"Public career 2020", "text/plain")},
        data={"ai_handling_policy": "local_only", "confirmed_public_information": "false"},
    )
    assert response.status_code == 422


def test_ingestion_history_corrections_reprocessing_and_suppression(client: TestClient) -> None:
    created = client.post(
        "/api/ingestions/documents",
        files={
            "file": (
                "history.txt",
                b"Public Professional\nBoard Chair 2021 - present",
                "text/plain",
            )
        },
        data={"ai_handling_policy": "local_only", "confirmed_public_information": "true"},
    ).json()
    created["proposal"]["profile"]["current_title"] = "Corrected title"
    corrected = client.put(f"/api/ingestions/{created['id']}/proposal", json=created["proposal"])
    assert corrected.status_code == 200
    assert corrected.json()["proposal"]["profile"]["current_title"] == "Corrected title"

    reprocessed = client.post(f"/api/ingestions/{created['id']}/reprocess")
    assert reprocessed.status_code == 200
    assert reprocessed.json()["provider"] == "deterministic"
    suppressed = client.post(f"/api/ingestions/{created['id']}/suppress")
    assert suppressed.status_code == 200
    assert suppressed.json()["status"] == "suppressed"
    assert client.get("/api/ingestions").json()[0]["id"] == created["id"]
    assert client.get("/api/ingestions/operations").status_code == 200


def test_asset_enrichment_adds_only_derived_fields(client: TestClient) -> None:
    asset = client.post(
        "/api/assets",
        json={
            "title": "Public research leadership",
            "description": "Led collaborative research partnerships and public engagement",
            "category": "Leadership Asset",
        },
    ).json()
    enriched = client.post(f"/api/ingestions/assets/{asset['id']}/enrich")
    assert enriched.status_code == 200
    updated = client.get(f"/api/assets/{asset['id']}").json()
    assert updated["title"] == asset["title"]
    assert updated["description"] == asset["description"]
    assert updated["tags"]


def test_multi_source_merge_deduplicates_and_preserves_provenance() -> None:
    institutional = PublicProfileSource(
        url="https://example.org/profile", source_type="institutional_profile"
    )
    scholar = PublicProfileSource(
        url="https://scholar.example.org/profile", source_type="google_scholar"
    )
    first = CareerExtractionProposal(
        profile=ProposedProfile(name="Example Professional", current_title="Director"),
        assets=[ProposedAsset(title="Research leadership", start_date="2020-01-01")],
    )
    second = CareerExtractionProposal(
        profile=ProposedProfile(name="Example Professional", current_title="Professor"),
        assets=[
            ProposedAsset(
                title="Research leadership",
                start_date="2020-01-01",
                themes=["Research impact"],
            )
        ],
    )
    merged = merge_source_proposals(
        [(institutional, "Institution", first), (scholar, "Scholar", second)]
    )
    assert len(merged.assets) == 1
    assert len(merged.assets[0].source_urls) == 2
    assert merged.assets[0].themes == ["Research impact"]
    assert merged.conflicts
    assert merged.coverage == {"institutional_profile": 1, "google_scholar": 1}
