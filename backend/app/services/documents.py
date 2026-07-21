import hashlib
import json
import re
from io import BytesIO
from pathlib import Path
from uuid import uuid4

from docx import Document as DocxDocument
from fastapi import HTTPException, UploadFile, status
from pypdf import PdfReader
from sqlmodel import Session, select

from app.core.config import get_settings
from app.models.career import AiHandlingPolicy, Document
from app.services.audit import record_audit

MAX_UPLOAD_BYTES = 20 * 1024 * 1024
ALLOWED_EXTENSIONS = {".pdf", ".docx", ".txt", ".json"}


def _safe_filename(filename: str) -> str:
    cleaned = re.sub(r"[^A-Za-z0-9._ -]", "_", Path(filename).name).strip(". ")
    return cleaned[:180] or "document"


def _extract_text(extension: str, content: bytes) -> str:
    if extension == ".pdf":
        return "\n\n".join(page.extract_text() or "" for page in PdfReader(BytesIO(content)).pages)
    if extension == ".docx":
        document = DocxDocument(BytesIO(content))
        blocks = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                cells = [re.sub(r"\s+", " ", cell.text).strip() for cell in row.cells]
                if any(cells):
                    blocks.append(" | ".join(cells))
        return "\n".join(blocks)
    decoded = content.decode("utf-8-sig")
    if extension == ".json":
        return json.dumps(json.loads(decoded), ensure_ascii=False, indent=2)
    return decoded


async def store_document(
    session: Session,
    *,
    upload: UploadFile,
    policy: AiHandlingPolicy,
    confirmed_public_information: bool,
) -> tuple[Document, bool]:
    if not confirmed_public_information:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_CONTENT,
            detail="Confirm that this document contains only public professional information",
        )
    filename = _safe_filename(upload.filename or "document")
    extension = Path(filename).suffix.lower()
    if extension not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=415, detail="Supported document types: PDF, DOCX, TXT, JSON"
        )
    content = await upload.read(MAX_UPLOAD_BYTES + 1)
    if not content:
        raise HTTPException(status_code=422, detail="The uploaded document is empty")
    if len(content) > MAX_UPLOAD_BYTES:
        raise HTTPException(status_code=413, detail="Documents must be 20 MB or smaller")

    digest = hashlib.sha256(content).hexdigest()
    existing = session.exec(select(Document).where(Document.sha256 == digest)).first()
    if existing:
        return existing, True

    document_id = uuid4()
    relative_path = Path("originals") / str(document_id) / filename
    absolute_path = get_settings().data_root / relative_path
    absolute_path.parent.mkdir(parents=True, exist_ok=True)
    temporary_path = absolute_path.with_suffix(f"{absolute_path.suffix}.tmp")
    temporary_path.write_bytes(content)
    temporary_path.replace(absolute_path)

    extraction_status = "completed"
    extraction_error = ""
    extracted_text = ""
    try:
        extracted_text = _extract_text(extension, content)
    except (OSError, UnicodeError, ValueError) as exc:
        extraction_status = "failed"
        extraction_error = str(exc)[:2_000]

    document = Document(
        id=document_id,
        original_filename=filename,
        mime_type=upload.content_type or "application/octet-stream",
        byte_size=len(content),
        sha256=digest,
        relative_path=relative_path.as_posix(),
        ai_handling_policy=policy.value,
        extracted_text=extracted_text,
        extraction_status=extraction_status,
        extraction_error=extraction_error,
    )
    session.add(document)
    record_audit(
        session,
        entity_type="document",
        entity_id=document.id,
        action="uploaded",
        details={"filename": filename, "sha256": digest, "policy": policy.value},
    )
    return document, False


def resolve_original_path(document: Document) -> Path:
    root = get_settings().data_root.resolve()
    path = (root / document.relative_path).resolve()
    if root not in path.parents:
        raise HTTPException(status_code=500, detail="Invalid stored document path")
    if not path.is_file():
        raise HTTPException(status_code=404, detail="Stored document file is missing")
    return path
