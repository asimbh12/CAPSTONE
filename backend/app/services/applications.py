import json
import logging
import re
from html import escape
from io import BytesIO
from typing import Any
from uuid import UUID

from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from fastapi import HTTPException
from reportlab.lib.colors import HexColor  # type: ignore[import-untyped]
from reportlab.lib.pagesizes import A4  # type: ignore[import-untyped]
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet  # type: ignore[import-untyped]
from reportlab.lib.units import mm  # type: ignore[import-untyped]
from reportlab.platypus import (  # type: ignore[import-untyped]
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
)
from sqlmodel import Session, col, select

from app.core.config import get_settings
from app.models.career import (
    ApplicationAssessment,
    ApplicationDraft,
    ApplicationRequirement,
    CareerAsset,
    CareerProfile,
    JobApplication,
)
from app.schemas.applications import (
    ApplicationRead,
    AssessmentRead,
    DraftRead,
    ProviderApplicationDrafts,
    RequirementInput,
    RequirementRead,
    RequirementsProposal,
)

STOPWORDS = {
    "about",
    "after",
    "also",
    "and",
    "are",
    "with",
    "from",
    "have",
    "into",
    "that",
    "the",
    "their",
    "this",
    "will",
    "your",
    "you",
    "for",
    "our",
    "role",
    "work",
    "must",
}
logger = logging.getLogger(__name__)


def get_or_404(session: Session, application_id: UUID) -> JobApplication:
    item = session.get(JobApplication, application_id)
    if item is None:
        raise HTTPException(status_code=404, detail="Job application not found")
    return item


def _words(text: str) -> set[str]:
    return {
        word for word in re.findall(r"[a-z][a-z0-9-]{2,}", text.casefold()) if word not in STOPWORDS
    }


def extract_requirements(text: str) -> list[RequirementInput]:
    settings = get_settings()
    if (
        settings.ai_provider.casefold() == "gemini"
        and settings.gemini_api_key
    ):
        try:
            from google import genai

            client = genai.Client(api_key=settings.gemini_api_key)
            response = client.models.generate_content(
                model=settings.gemini_model,
                contents=(
                    "Extract every explicit essential and desirable requirement from this "
                    "position description. Separate compound criteria when they require "
                    "distinct evidence. Preserve the source meaning, classify each item, and "
                    "assign weight 2 to explicit essential or mandatory criteria and 1 to "
                    "desirable criteria. Do not invent requirements.\n\n"
                    f"{text[:120_000]}"
                ),
                config={
                    "response_mime_type": "application/json",
                    "response_schema": RequirementsProposal,
                    "temperature": 0,
                },
            )
            return RequirementsProposal.model_validate_json(
                response.text or "{}"
            ).requirements
        except Exception:
            # The local extractor keeps intake available when Gemini is unavailable.
            pass
    lines = [re.sub(r"\s+", " ", line).strip(" •\t-") for line in text.splitlines()]
    lines = [line for line in lines if 15 <= len(line) <= 800]
    headings = re.compile(r"selection criteria|requirements|essential|desirable|capabilities", re.I)
    start = next((index for index, line in enumerate(lines) if headings.search(line)), 0)
    candidates = lines[start + 1 :] if start else lines
    signals = re.compile(
        r"\b(ability|experience|qualification|knowledge|skill|demonstrated|track record|"
        r"leadership|communication|research|management|must|required|degree)\b",
        re.I,
    )
    selected = [line for line in candidates if signals.search(line)]
    if not selected:
        selected = candidates[:12]
    seen: set[str] = set()
    result: list[RequirementInput] = []
    for line in selected:
        title = re.sub(
            r"^(essential|desirable|criterion|criteria)\s*[:.-]?\s*", "", line, flags=re.I
        )
        title = title[:300]
        key = re.sub(r"\W+", "", title.casefold())
        if not key or key in seen:
            continue
        seen.add(key)
        requirement_type = "desirable" if "desirable" in line.casefold() else "essential"
        result.append(
            RequirementInput(title=title, description=line, requirement_type=requirement_type)
        )
        if len(result) == 20:
            break
    if not result:
        result.append(
            RequirementInput(title="Review the position description", description=text[:1000])
        )
    return result


def _requirement_read(row: ApplicationRequirement) -> RequirementRead:
    return RequirementRead(
        **row.model_dump(exclude={"asset_ids_json", "application_id", "created_at"}),
        asset_ids=json.loads(row.asset_ids_json),
    )


def build_read(session: Session, item: JobApplication) -> ApplicationRead:
    requirements = session.exec(
        select(ApplicationRequirement)
        .where(ApplicationRequirement.application_id == item.id)
        .order_by(col(ApplicationRequirement.sort_order))
    ).all()
    assessment = session.exec(
        select(ApplicationAssessment)
        .where(ApplicationAssessment.application_id == item.id)
        .order_by(col(ApplicationAssessment.version).desc())
    ).first()
    drafts = session.exec(
        select(ApplicationDraft)
        .where(ApplicationDraft.application_id == item.id)
        .order_by(col(ApplicationDraft.created_at).desc())
    ).all()
    assessment_read = None
    if assessment:
        assessment_read = AssessmentRead(
            **assessment.model_dump(
                exclude={"strengths_json", "gaps_json", "recommendations_json", "application_id"}
            ),
            strengths=json.loads(assessment.strengths_json),
            gaps=json.loads(assessment.gaps_json),
            recommendations=json.loads(assessment.recommendations_json),
        )
    return ApplicationRead(
        **item.model_dump(),
        requirements=[_requirement_read(row) for row in requirements],
        assessment=assessment_read,
        drafts=[
            DraftRead(
                **draft.model_dump(exclude={"application_id", "unsupported_claims_json"}),
                unsupported_claims=json.loads(draft.unsupported_claims_json),
            )
            for draft in drafts
        ],
    )


def map_and_assess(session: Session, item: JobApplication) -> ApplicationAssessment:
    if not item.requirements_confirmed:
        raise HTTPException(
            status_code=409, detail="Review and confirm the extracted requirements first"
        )
    requirements = list(
        session.exec(
            select(ApplicationRequirement).where(ApplicationRequirement.application_id == item.id)
        ).all()
    )
    assets = list(session.exec(select(CareerAsset).where(CareerAsset.status == "active")).all())
    for requirement in requirements:
        requirement_words = _words(f"{requirement.title} {requirement.description}")
        ranked: list[tuple[float, CareerAsset]] = []
        for asset in assets:
            asset_words = _words(
                f"{asset.title} {asset.description} {asset.impact_summary} "
                f"{asset.role} {asset.tags_json} {asset.keywords_json}"
            )
            overlap = len(requirement_words & asset_words)
            score = overlap / max(1, min(len(requirement_words), 12))
            if overlap:
                ranked.append((score, asset))
        ranked.sort(key=lambda value: value[0], reverse=True)
        matched = [asset for score, asset in ranked[:5] if score >= 0.08]
        best = ranked[0][0] if ranked else 0
        requirement.asset_ids_json = json.dumps([str(asset.id) for asset in matched])
        requirement.coverage = round(min(100, best * 220 + min(len(matched), 3) * 12), 1)
        requirement.confidence = (
            round(min(100, 40 + best * 100 + len(matched) * 8), 1) if matched else 20
        )
        requirement.explanation = (
            f"Matched {len(matched)} existing career asset(s) using shared "
            "requirement and evidence terms."
            if matched
            else "No sufficiently related career asset was found; treat this as a "
            "gap until evidence is added."
        )
        session.add(requirement)
    total_weight = sum(row.weight for row in requirements) or 1
    fit = sum(row.coverage * row.weight for row in requirements) / total_weight
    strengths = [row.title for row in requirements if row.coverage >= 65]
    gaps = [row.title for row in requirements if row.coverage < 40]
    previous = session.exec(
        select(ApplicationAssessment).where(ApplicationAssessment.application_id == item.id)
    ).all()
    assessment = ApplicationAssessment(
        application_id=item.id,
        version=len(previous) + 1,
        fit_score=round(fit, 1),
        overall_confidence=round(
            sum(row.confidence for row in requirements) / max(1, len(requirements)), 1
        ),
        strengths_json=json.dumps(strengths),
        gaps_json=json.dumps(gaps),
        recommendations_json=json.dumps(
            [f"Add or strengthen evidence for: {gap}" for gap in gaps[:5]]
        ),
    )
    session.add(assessment)
    return assessment


def _asset_context(assets: list[CareerAsset]) -> str:
    return json.dumps(
        [
            {
                "id": str(asset.id),
                "title": asset.title,
                "role": asset.role,
                "category": asset.category,
                "description": asset.description,
                "impact": asset.impact_summary,
                "start_date": str(asset.start_date) if asset.start_date else None,
                "end_date": str(asset.end_date) if asset.end_date else None,
                "tags": json.loads(asset.tags_json),
                "keywords": json.loads(asset.keywords_json),
            }
            for asset in assets
        ],
        ensure_ascii=False,
    )


def _provider_drafts(
    item: JobApplication,
    profile: CareerProfile | None,
    requirements: list[ApplicationRequirement],
    assets: list[CareerAsset],
) -> ProviderApplicationDrafts | None:
    settings = get_settings()
    if settings.ai_provider.casefold() != "gemini" or not settings.gemini_api_key:
        return None
    from google import genai

    profile_context = {
        "name": profile.name if profile else "",
        "current_title": profile.current_title if profile else "",
        "current_organisation": profile.current_organisation if profile else "",
        "career_mission": profile.career_mission if profile else "",
        "career_narrative": profile.career_narrative if profile else "",
    }
    requirement_context = [
        {
            "title": row.title,
            "description": row.description,
            "type": row.requirement_type,
            "weight": row.weight,
            "mapped_asset_ids": json.loads(row.asset_ids_json),
        }
        for row in requirements
    ]
    prompt = f"""
Prepare a senior-executive application pack using only the supplied public career facts. Never
invent responsibilities, outcomes, dates, qualifications, metrics, employers or awards. Record
any important position requirement that lacks evidence in unsupported_claims instead of filling it.

Produce four genuinely useful, distinct documents in clean Markdown:
- cover_letter: 800-1200 words; role-specific persuasive narrative; 3-5 evidence-led themes;
  explicit organisational contribution; confident close. Do not produce an evidence shopping list.
- selection_criteria: every reviewed criterion under its own ## heading; detailed context, actions,
  outcomes and relevance where supported; honest limitations; varied examples without repetition.
- tailored_cv: executive profile, core capabilities, selected leadership impact, relevant career
  experience, research/technology leadership, stakeholder/government/industry engagement, awards
  and qualifications only when supported. Preserve useful verified metrics.
- interview_notes: executive value proposition, 6-10 likely questions, evidence-backed talking
  points, 4-6 STAR story outlines, risks with honest bridging language, panel questions and a
  proposed 90-day contribution outline clearly distinguished from past achievements.

Use ## and ### headings, normal paragraphs, and '- ' bullets. Do not use tables. Tailor every
document to the role. The four outputs must not be repetitions of the same asset list.

ROLE: {item.role_title}
ORGANISATION: {item.organisation}
POSITION DESCRIPTION:
{item.position_description[:100_000]}

PROFILE: {json.dumps(profile_context, ensure_ascii=False)}
REVIEWED REQUIREMENTS: {json.dumps(requirement_context, ensure_ascii=False)}
VERIFIED CAREER ASSETS: {_asset_context(assets)[:120_000]}
"""
    try:
        client = genai.Client(api_key=settings.gemini_api_key)
        response = client.models.generate_content(
            model=settings.gemini_model,
            contents=prompt,
            config={
                "response_mime_type": "application/json",
                "response_schema": ProviderApplicationDrafts,
                "temperature": 0.2,
            },
        )
        return ProviderApplicationDrafts.model_validate_json(response.text or "{}")
    except Exception:
        logger.exception("Gemini application drafting failed; using grounded local fallback")
        return None


def generate_drafts(session: Session, item: JobApplication) -> list[ApplicationDraft]:
    assessment = session.exec(
        select(ApplicationAssessment)
        .where(ApplicationAssessment.application_id == item.id)
        .order_by(col(ApplicationAssessment.version).desc())
    ).first()
    if assessment is None:
        raise HTTPException(
            status_code=409, detail="Run the evidence fit assessment before generating drafts"
        )
    requirements = list(
        session.exec(
            select(ApplicationRequirement)
            .where(ApplicationRequirement.application_id == item.id)
            .order_by(col(ApplicationRequirement.sort_order))
        ).all()
    )
    asset_ids = {UUID(value) for row in requirements for value in json.loads(row.asset_ids_json)}
    assets = (
        {
            asset.id: asset
            for asset in session.exec(
                select(CareerAsset).where(col(CareerAsset.id).in_(asset_ids))
            ).all()
        }
        if asset_ids
        else {}
    )
    profile = session.exec(select(CareerProfile)).first()
    name = profile.name if profile and profile.name else "Applicant"
    evidence_lines = [
        f"• {asset.title}: {asset.impact_summary or asset.description}" for asset in assets.values()
    ]
    evidence = "\n".join(evidence_lines) or "• No mapped evidence is currently available."
    criteria = "\n\n".join(
        f"{row.title}\n"
        + (
            "\n".join(
                f"• {assets[UUID(asset_id)].title}: "
                f"{assets[UUID(asset_id)].impact_summary or assets[UUID(asset_id)].description}"
                for asset_id in json.loads(row.asset_ids_json)
                if UUID(asset_id) in assets
            )
            or "Evidence gap: add a verified example before submission."
        )
        for row in requirements
    )
    contents = {
        "cover_letter": (
            "Dear Hiring Committee,\n\n"
            f"I am applying for the {item.role_title} position at "
            f"{item.organisation or 'your organisation'}. My relevant, verified "
            f"career evidence includes:\n\n{evidence}\n\nI would welcome the opportunity "
            "to discuss how this experience aligns with the role.\n\n"
            f"Yours sincerely,\n{name}"
        ),
        "selection_criteria": criteria,
        "tailored_cv": (
            f"{name}\nTarget role: {item.role_title}\n\n"
            f"RELEVANT CAREER EVIDENCE\n{evidence}"
        ),
        "interview_notes": (
            f"INTERVIEW PREPARATION — {item.role_title}\n\n"
            f"Evidence to foreground:\n{evidence}\n\n"
            "Evidence gaps to prepare honestly:\n"
        )
        + "\n".join(f"• {gap}" for gap in json.loads(assessment.gaps_json)),
    }
    all_assets = list(
        session.exec(
            select(CareerAsset)
            .where(CareerAsset.status == "active")
            .order_by(col(CareerAsset.start_date).desc(), col(CareerAsset.created_at).desc())
        ).all()
    )
    provider_result = _provider_drafts(item, profile, requirements, all_assets)
    provider = "grounded_template"
    unsupported_claims: list[str] = []
    if provider_result:
        contents = provider_result.model_dump(exclude={"unsupported_claims"})
        provider = "gemini"
        unsupported_claims = provider_result.unsupported_claims
    for existing in session.exec(
        select(ApplicationDraft).where(ApplicationDraft.application_id == item.id)
    ).all():
        session.delete(existing)
    drafts = []
    for draft_type, content in contents.items():
        draft = ApplicationDraft(
            application_id=item.id,
            draft_type=draft_type,
            content=content.strip(),
            provider=provider,
            unsupported_claims_json=json.dumps(unsupported_claims),
        )
        session.add(draft)
        drafts.append(draft)
    return drafts


def _markdown_text(value: str) -> str:
    return value.replace("*", "").replace("__", "").replace("`", "").strip()


def _add_docx_markdown(document: Any, content: str) -> None:
    for line in content.splitlines():
        text = line.strip()
        if not text:
            continue
        if text.startswith("### "):
            document.add_heading(_markdown_text(text[4:]), level=2)
        elif text.startswith("## "):
            document.add_heading(_markdown_text(text[3:]), level=1)
        elif text.startswith("# "):
            document.add_heading(_markdown_text(text[2:]), level=1)
        elif text.startswith("- "):
            document.add_paragraph(_markdown_text(text[2:]), style="List Bullet")
        else:
            document.add_paragraph(_markdown_text(text))


def _add_pdf_markdown(
    story: list[object],
    content: str,
    styles: dict[str, ParagraphStyle],
) -> None:
    for line in content.splitlines():
        text = line.strip()
        if not text:
            story.append(Spacer(1, 4))
        elif text.startswith("### "):
            story.append(Paragraph(escape(_markdown_text(text[4:])), styles["heading3"]))
        elif text.startswith(("## ", "# ")):
            story.append(
                Paragraph(escape(_markdown_text(text.lstrip("# "))), styles["heading2"])
            )
        elif text.startswith("- "):
            story.append(
                Paragraph(f"- {escape(_markdown_text(text[2:]))}", styles["bullet"])
            )
        else:
            story.append(Paragraph(escape(_markdown_text(text)), styles["body"]))


def export_pack(session: Session, item: JobApplication, format_name: str) -> bytes:
    latest_by_type: dict[str, ApplicationDraft] = {}
    for draft in session.exec(
        select(ApplicationDraft)
        .where(ApplicationDraft.application_id == item.id)
        .order_by(col(ApplicationDraft.created_at).desc())
    ).all():
        latest_by_type.setdefault(draft.draft_type, draft)
    if not latest_by_type:
        raise HTTPException(status_code=409, detail="Generate application drafts before exporting")
    titles = {
        "cover_letter": "Cover letter",
        "selection_criteria": "Selection criteria",
        "tailored_cv": "Tailored CV",
        "interview_notes": "Interview notes",
    }
    if format_name == "docx":
        document = DocxDocument()
        section = document.sections[0]
        section.top_margin = section.bottom_margin = Inches(1)
        section.left_margin = section.right_margin = Inches(1)
        styles = document.styles
        styles["Normal"].font.name = "Calibri"
        styles["Normal"].font.size = Pt(11)
        styles["Normal"].paragraph_format.space_after = Pt(8)
        styles["Normal"].paragraph_format.line_spacing = 1.15
        for style_name, size, before, after in (
            ("Heading 1", 16, 18, 10),
            ("Heading 2", 13, 12, 6),
            ("Heading 3", 12, 8, 4),
        ):
            style = styles[style_name]
            style.font.name = "Calibri"
            style.font.size = Pt(size)
            style.font.color.rgb = RGBColor(46, 116, 181)
            style.paragraph_format.space_before = Pt(before)
            style.paragraph_format.space_after = Pt(after)
        title = document.add_heading(item.role_title, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.runs[0].font.color.rgb = RGBColor(31, 78, 121)
        document.add_paragraph(
            item.organisation, style="Subtitle"
        ).alignment = WD_ALIGN_PARAGRAPH.CENTER
        for index, draft_type in enumerate(titles):
            if draft_type not in latest_by_type:
                continue
            if index:
                document.add_page_break()  # type: ignore[no-untyped-call]
            document.add_heading(titles[draft_type], level=1)
            _add_docx_markdown(document, latest_by_type[draft_type].content)
        buffer = BytesIO()
        document.save(buffer)
        return buffer.getvalue()
    buffer = BytesIO()
    pdf = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=18 * mm,
        rightMargin=18 * mm,
        topMargin=16 * mm,
        bottomMargin=16 * mm,
        title=f"{item.role_title} application pack",
    )
    styles = getSampleStyleSheet()
    export_styles = {
        "heading2": ParagraphStyle(
            "ApplicationHeading2",
            parent=styles["Heading2"],
            textColor=HexColor("#2E74B5"),
            fontSize=13,
            leading=16,
            spaceBefore=12,
            spaceAfter=6,
        ),
        "heading3": ParagraphStyle(
            "ApplicationHeading3",
            parent=styles["Heading3"],
            textColor=HexColor("#1F4D78"),
            fontSize=11.5,
            leading=14,
            spaceBefore=8,
            spaceAfter=4,
        ),
        "body": ParagraphStyle(
            "ApplicationBody",
            parent=styles["BodyText"],
            fontName="Helvetica",
            fontSize=10.5,
            leading=14,
            spaceAfter=7,
        ),
        "bullet": ParagraphStyle(
            "ApplicationBullet",
            parent=styles["BodyText"],
            fontName="Helvetica",
            fontSize=10.5,
            leading=14,
            leftIndent=16,
            firstLineIndent=-10,
            spaceAfter=5,
        ),
    }
    heading = ParagraphStyle(
        "CapstoneHeading",
        parent=styles["Heading1"],
        textColor=HexColor("#1F4E79"),
        spaceAfter=10,
    )
    story = [
        Paragraph(item.role_title, styles["Title"]),
        Paragraph(item.organisation or "Application pack", styles["Heading2"]),
        Spacer(1, 10),
    ]
    for index, draft_type in enumerate(titles):
        if draft_type not in latest_by_type:
            continue
        if index:
            story.append(PageBreak())
        story.append(Paragraph(titles[draft_type], heading))
        _add_pdf_markdown(story, latest_by_type[draft_type].content, export_styles)
    pdf.build(story)
    return buffer.getvalue()
