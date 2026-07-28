import json
import logging
from html import escape
from io import BytesIO
from typing import Any
from uuid import UUID

from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from fastapi import HTTPException
from reportlab.lib.colors import HexColor  # type: ignore[import-untyped]
from reportlab.lib.pagesizes import A4  # type: ignore[import-untyped]
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet  # type: ignore[import-untyped]
from reportlab.lib.units import mm  # type: ignore[import-untyped]
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer  # type: ignore[import-untyped]
from sqlmodel import Session, col, select

from app.core.config import get_settings
from app.models.career import AiOperation, CareerAsset, CareerDocument, CareerProfile
from app.schemas.career_documents import (
    CareerDocumentGenerate,
    CareerDocumentRead,
    ProviderCareerDocument,
)
from app.services.applications import _add_docx_markdown, _add_pdf_markdown, _asset_context

logger = logging.getLogger(__name__)

TYPE_LABELS = {
    "professional_biography": "Professional biography",
    "executive_profile": "Executive profile",
    "linkedin_about": "LinkedIn About",
    "academic_cv": "Academic CV",
    "executive_cv": "Executive CV",
    "board_cv": "Board CV",
    "grant_cv": "Two-page grant CV",
    "capability_statement": "Capability statement",
}

TYPE_GUIDANCE = {
    "professional_biography": (
        "Write 350-500 words with a strong opening identity, career focus, selected leadership "
        "and impact, and a concise closing."
    ),
    "executive_profile": (
        "Write 600-900 words with an executive value proposition, leadership scope, selected "
        "impact themes and future-facing contribution."
    ),
    "linkedin_about": (
        "Write 250-400 words in first person with a distinctive opening, accessible impact "
        "narrative, focus areas and a professional closing. Avoid generic buzzwords."
    ),
    "academic_cv": (
        "Create a detailed academic CV with profile, appointments, research leadership, grants, "
        "publications, supervision, teaching, service, awards and qualifications only where "
        "supported. Use reverse chronology and preserve verified dates and metrics."
    ),
    "executive_cv": (
        "Create a focused executive CV with executive profile, core capabilities, leadership "
        "experience, enterprise impact, stakeholder engagement, governance, recognition and "
        "qualifications only where supported. Prioritise outcomes over duties."
    ),
    "board_cv": (
        "Create a board CV with governance proposition, board and committee experience, "
        "strategy, risk, stakeholder, financial or organisational oversight evidence, sector "
        "expertise and recognition. Do not imply director duties not present in evidence."
    ),
    "grant_cv": (
        "Create a concise two-page grant CV with research profile, selected track record, "
        "funding and translation, leadership, team capability, selected outputs and contribution "
        "to the proposed purpose. Be selective and avoid unsupported totals."
    ),
    "capability_statement": (
        "Create a persuasive capability statement with purpose, distinctive capabilities, "
        "evidence of delivery, leadership and partnerships, selected proof points and a concise "
        "value proposition tailored to the audience."
    ),
}


def read_document(item: CareerDocument) -> CareerDocumentRead:
    return CareerDocumentRead(
        **item.model_dump(exclude={"asset_ids_json", "unsupported_claims_json"}),
        asset_ids=json.loads(item.asset_ids_json),
        unsupported_claims=json.loads(item.unsupported_claims_json),
    )


def get_or_404(session: Session, document_id: UUID) -> CareerDocument:
    item = session.get(CareerDocument, document_id)
    if item is None:
        raise HTTPException(status_code=404, detail="Career document not found")
    return item


def _fallback(
    payload: CareerDocumentGenerate,
    profile: CareerProfile | None,
    assets: list[CareerAsset],
) -> str:
    name = profile.name if profile and profile.name else "Career professional"
    current = " at ".join(
        value
        for value in (
            profile.current_title if profile else "",
            profile.current_organisation if profile else "",
        )
        if value
    )
    narrative = (
        profile.career_narrative or profile.career_mission
        if profile
        else ""
    )
    evidence = [asset.impact_summary or asset.description or asset.title for asset in assets]
    if payload.document_type == "linkedin_about":
        return "\n\n".join(
            part
            for part in (
                f"I am {name}{f', {current}' if current else ''}.",
                narrative,
                "My work includes:\n" + "\n".join(f"- {value}" for value in evidence),
                payload.purpose,
            )
            if part
        )
    heading = TYPE_LABELS[payload.document_type]
    if payload.document_type in {
        "academic_cv",
        "executive_cv",
        "board_cv",
        "grant_cv",
        "capability_statement",
    }:
        grouped: dict[str, list[str]] = {}
        for asset in assets:
            date_label = str(asset.start_date.year) if asset.start_date else "Date not recorded"
            detail = asset.impact_summary or asset.description or asset.title
            grouped.setdefault(asset.category, []).append(
                f"- **{asset.title}** ({date_label}) — {detail}"
            )
        sections = "\n\n".join(
            f"## {category}\n" + "\n".join(entries)
            for category, entries in grouped.items()
        )
        return "\n\n".join(
            part
            for part in (
                f"# {heading}",
                f"## Professional profile\n{name}{f' — {current}' if current else ''}. "
                f"{narrative}".strip(),
                f"## Purpose\n{payload.purpose}" if payload.purpose else "",
                sections,
                "## Evidence note\nThis draft contains only selected active CAPSTONE assets.",
            )
            if part
        )
    return "\n\n".join(
        part
        for part in (
            f"# {heading}",
            f"{name}{f' is {current}' if current else ''}. {narrative}".strip(),
            "## Selected impact\n" + "\n".join(f"- {value}" for value in evidence),
            f"## Intended use\n{payload.purpose}" if payload.purpose else "",
        )
        if part
    )


def _provider_content(
    payload: CareerDocumentGenerate,
    profile: CareerProfile | None,
    assets: list[CareerAsset],
) -> tuple[ProviderCareerDocument | None, str]:
    settings = get_settings()
    if settings.ai_provider.casefold() != "gemini" or not settings.gemini_api_key:
        return None, ""
    profile_data = profile.model_dump(mode="json") if profile else {}
    audience = payload.audience or "general professional audience"
    purpose = payload.purpose or "reusable career communication"
    prompt = f"""
Write a polished {TYPE_LABELS[payload.document_type]} using only the verified public facts below.
Never invent roles, dates, qualifications, awards, metrics, outcomes or endorsements. Put any
requested idea that lacks evidence into unsupported_claims instead of stating it. Use clean
Markdown headings, paragraphs and bullets; no tables. Audience: {audience}. Purpose: {purpose}.
Tone: {payload.tone}. Make the document specific, cohesive and impact-led rather than an asset
list.

FORMAT AND DEPTH: {TYPE_GUIDANCE[payload.document_type]}

PROFILE: {json.dumps(profile_data, ensure_ascii=False)}
VERIFIED ASSETS: {_asset_context(assets)[:120_000]}
"""
    try:
        from google import genai

        client = genai.Client(api_key=settings.gemini_api_key)
        response = client.models.generate_content(
            model=settings.gemini_model,
            contents=prompt,
            config={
                "response_mime_type": "application/json",
                "response_schema": ProviderCareerDocument,
                "temperature": 0.25,
            },
        )
        return ProviderCareerDocument.model_validate_json(response.text or "{}"), ""
    except Exception as exc:
        logger.exception("Gemini career-document generation failed; using local fallback")
        message = f"{type(exc).__name__}: {exc}"
        if settings.gemini_api_key:
            message = message.replace(settings.gemini_api_key, "[redacted]")
        return None, message[:2_000]


def generate(session: Session, payload: CareerDocumentGenerate) -> CareerDocument:
    profile = session.exec(select(CareerProfile)).first()
    query = select(CareerAsset).where(CareerAsset.status == "active")
    if payload.asset_ids:
        query = query.where(col(CareerAsset.id).in_(payload.asset_ids))
    assets = list(
        session.exec(
            query.order_by(col(CareerAsset.start_date).desc(), col(CareerAsset.created_at).desc())
        ).all()
    )
    if not assets:
        raise HTTPException(
            status_code=409,
            detail="Add at least one active career asset before generating a document",
        )
    provider_result, provider_error = _provider_content(payload, profile, assets)
    item = CareerDocument(
        document_type=payload.document_type,
        title=payload.title,
        audience=payload.audience,
        purpose=payload.purpose,
        tone=payload.tone,
        content=provider_result.content if provider_result else _fallback(payload, profile, assets),
        provider="gemini" if provider_result else "grounded_template",
        asset_ids_json=json.dumps([str(asset.id) for asset in assets]),
        unsupported_claims_json=json.dumps(
            provider_result.unsupported_claims if provider_result else []
        ),
    )
    session.add(item)
    settings = get_settings()
    if settings.ai_provider.casefold() == "gemini" and settings.gemini_api_key:
        session.add(
            AiOperation(
                operation="generate_career_document",
                entity_type="career_document",
                entity_id=str(item.id),
                provider="gemini",
                model=settings.gemini_model,
                status="succeeded" if provider_result else "failed",
                input_characters=len(_asset_context(assets)) + len(payload.purpose),
                output_characters=len(item.content),
                error_message=provider_error,
            )
        )
    return item


def export_document(item: CareerDocument, format_name: str) -> bytes:
    if format_name == "docx":
        document = DocxDocument()
        section = document.sections[0]
        section.top_margin = section.bottom_margin = Inches(0.8)
        section.left_margin = section.right_margin = Inches(0.9)
        document.styles["Normal"].font.name = "Calibri"
        document.styles["Normal"].font.size = Pt(11)
        title = document.add_heading(item.title, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.runs[0].font.color.rgb = RGBColor(31, 78, 121)
        if item.audience:
            subtitle = document.add_paragraph(item.audience, style="Subtitle")
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_docx_markdown(document, item.content)
        buffer = BytesIO()
        document.save(buffer)
        return buffer.getvalue()
    buffer = BytesIO()
    pdf = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=18 * mm,
        rightMargin=18 * mm,
        topMargin=16 * mm,
        bottomMargin=16 * mm,
        title=item.title,
    )
    sample = getSampleStyleSheet()
    styles: dict[str, ParagraphStyle] = {
        "heading2": ParagraphStyle(
            "CareerHeading2", parent=sample["Heading2"], textColor=HexColor("#2E74B5")
        ),
        "heading3": ParagraphStyle(
            "CareerHeading3", parent=sample["Heading3"], textColor=HexColor("#1F4D78")
        ),
        "body": ParagraphStyle(
            "CareerBody", parent=sample["BodyText"], fontSize=10.5, leading=14, spaceAfter=7
        ),
        "bullet": ParagraphStyle(
            "CareerBullet", parent=sample["BodyText"], leftIndent=16, firstLineIndent=-10
        ),
        "subtitle": ParagraphStyle(
            "CareerSubtitle",
            parent=sample["BodyText"],
            alignment=1,
            textColor=HexColor("#5B6573"),
            fontSize=11,
            leading=14,
        ),
    }
    story: list[Any] = [
        Paragraph(escape(item.title), sample["Title"]),
        Paragraph(escape(item.audience), styles["subtitle"])
        if item.audience
        else Spacer(1, 4),
        Spacer(1, 10),
    ]
    _add_pdf_markdown(story, item.content, styles)
    pdf.build(story)
    return buffer.getvalue()
